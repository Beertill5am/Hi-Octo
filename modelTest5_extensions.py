import os
import re
import time
from typing import Protocol, Type, Dict, List, Any, Callable, TypeVar
from pydantic import BaseModel
from docx import Document as DocxDocument
import fitz 


# =============================================================================
# SECTION 1: DOCUMENT SPECIALISTS
# =============================================================================

class DocumentSpecialist(Protocol):
    """Interface that all specialists must implement."""
    def convert(self, file_path: str) -> str:
        ...


class TXTSpecialist:
    """Handler for plain text files."""
    
    def _wrap_metadata(self, content: str, source: str) -> str:
        filename = os.path.basename(source)
        return f"""---
            title: "{filename}"
            source: "{source}"
            type: "txt"
            ---

            {content}
            """
    
    def convert(self, file_path: str) -> str:
        print(f"📄 TXT: Reading '{os.path.basename(file_path)}'...")
        try:
            # Try multiple encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return f"❌ TXT Error: Unable to decode file with supported encodings"
            
            return self._wrap_metadata(content, file_path)
        except Exception as e:
            return f"❌ TXT Error: {e}"


class MarkdownSpecialist:
    """Handler for Markdown files - passthrough with metadata."""
    
    def _wrap_metadata(self, content: str, source: str) -> str:
        filename = os.path.basename(source)
        return f"""---
            title: "{filename}"
            source: "{source}"
            type: "md"
            ---

            {content}
            """
    
    def convert(self, file_path: str) -> str:
        print(f"📝 Markdown: Reading '{os.path.basename(file_path)}'...")
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Check if file already has YAML frontmatter
            if content.strip().startswith('---'):
                return content  # Already has metadata, pass through
            
            return self._wrap_metadata(content, file_path)
        except Exception as e:
            return f"❌ Markdown Error: {e}"


class DOCXSpecialist:
    """Handler for Microsoft Word documents."""
    
    def _wrap_metadata(self, content: str, title: str, source: str) -> str:
        return f"""---
            title: "{title}"
            source: "{source}"
            type: "docx"
            ---

            {content}
            """
    
    def convert(self, file_path: str) -> str:
        print(f"📘 DOCX: Converting '{os.path.basename(file_path)}'...")
        try:
            doc = DocxDocument(file_path)
            title = doc.core_properties.title or os.path.basename(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect heading styles and convert to Markdown
                style_name = para.style.name if para.style else ""
                if style_name.startswith('Heading'):
                    # Extract heading level (Heading 1, Heading 2, etc.)
                    level_match = re.search(r'\d+', style_name)
                    level = int(level_match.group()) if level_match else 1
                    paragraphs.append('#' * level + ' ' + text)
                elif style_name == 'Title':
                    paragraphs.append('# ' + text)
                else:
                    paragraphs.append(text)
            
            # Also extract text from tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append('| ' + ' | '.join(cells) + ' |')
                if table_rows:
                    # Add header separator after first row
                    header_sep = '| ' + ' | '.join(['---'] * len(table.rows[0].cells)) + ' |'
                    table_rows.insert(1, header_sep)
                    paragraphs.append('\n'.join(table_rows))
            
            content = '\n\n'.join(paragraphs)
            return self._wrap_metadata(content, title, file_path)
            
        except Exception as e:
            return f"❌ DOCX Error: {e}"


class PDFSpecialist:
    """Handler for PDF documents using PyMuPDF."""
    
    def _wrap_metadata(self, content: str, title: str, author: str, 
                       source: str, page_count: int) -> str:
        return f"""---
            title: "{title}"
            author: "{author}"
            source: "{source}"
            type: "pdf"
            pages: {page_count}
            ---

            {content}
            """
    
    def convert(self, file_path: str) -> str:
        if not PDF_AVAILABLE:
            return "❌ PDF Error: PyMuPDF not installed. Run: pip install pymupdf"
        
        print(f"📕 PDF: Converting '{os.path.basename(file_path)}'...")
        try:
            doc = fitz.open(file_path)
            
            # Extract metadata
            metadata = doc.metadata or {}
            title = metadata.get('title') or os.path.basename(file_path)
            author = metadata.get('author') or 'Unknown'
            
            # Clean up title/author (remove null bytes, etc.)
            title = title.replace('\x00', '').strip() or os.path.basename(file_path)
            author = author.replace('\x00', '').strip() or 'Unknown'
            
            pages = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                
                if text.strip():
                    # Add page marker as header
                    pages.append(f"## Page {page_num + 1}\n\n{text.strip()}")
            
            content = '\n\n---\n\n'.join(pages)
            page_count = len(doc)
            doc.close()
            
            return self._wrap_metadata(content, title, author, file_path, page_count)
            
        except Exception as e:
            return f"❌ PDF Error: {e}"


def register_all_specialists(router):
    """
    Register all document specialists with the ingestion router.
    
    Usage:
        from modelTest5_extensions import register_all_specialists
        register_all_specialists(router)
    """
    router.register(".txt", TXTSpecialist)
    router.register(".md", MarkdownSpecialist)
    router.register(".docx", DOCXSpecialist)
    router.register(".pdf", PDFSpecialist)
    
    print("✅ All available document specialists registered.")


# =============================================================================
# SECTION 2: RESOURCE-AWARE PRE-FLIGHT ANALYSIS
# =============================================================================

REFINEMENT_THRESHOLDS = {
    "max_chunks": 50,      
    "max_words": 10000,   
    "max_pages": 20        
}


def preflight_analysis(splits: List[Any]) -> Dict[str, Any]:
    """
    Pattern: Resource-Aware Optimization
    Analyzes document chunks BEFORE sending to LLM for refinement.
    
    Args:
        splits: List of Document objects from chunking
        
    Returns:
        Dict with metrics and skip_refinement recommendation
    """
    total_chunks = len(splits)
    total_words = sum(len(doc.page_content.split()) for doc in splits)
    
    # Estimate page count from metadata if available
    pages = set()
    for doc in splits:
        if hasattr(doc, 'metadata') and 'page' in doc.metadata:
            pages.add(doc.metadata['page'])
        elif hasattr(doc, 'metadata') and 'pages' in doc.metadata:
            pages.add(doc.metadata['pages'])
    
    # Fallback: estimate ~3 chunks per page
    estimated_pages = len(pages) if pages else max(1, total_chunks // 3)
    
    # Decision logic
    skip_refinement = (
        total_chunks > REFINEMENT_THRESHOLDS["max_chunks"] or
        total_words > REFINEMENT_THRESHOLDS["max_words"] or
        estimated_pages > REFINEMENT_THRESHOLDS["max_pages"]
    )
    
    analysis = {
        "total_chunks": total_chunks,
        "total_words": total_words,
        "estimated_pages": estimated_pages,
        "skip_refinement": skip_refinement,
        "reason": None
    }
    
    # Determine the specific reason for skipping
    if skip_refinement:
        if total_chunks > REFINEMENT_THRESHOLDS["max_chunks"]:
            analysis["reason"] = f"Chunks ({total_chunks}) > threshold ({REFINEMENT_THRESHOLDS['max_chunks']})"
        elif total_words > REFINEMENT_THRESHOLDS["max_words"]:
            analysis["reason"] = f"Words ({total_words:,}) > threshold ({REFINEMENT_THRESHOLDS['max_words']:,})"
        else:
            analysis["reason"] = f"Pages ({estimated_pages}) > threshold ({REFINEMENT_THRESHOLDS['max_pages']})"
    
    return analysis


def print_preflight_report(analysis: Dict[str, Any]) -> None:
    """Print a formatted pre-flight analysis report."""
    print("\n📊 PRE-FLIGHT ANALYSIS")
    print("=" * 40)
    print(f"   Chunks: {analysis['total_chunks']}")
    print(f"   Words:  {analysis['total_words']:,}")
    print(f"   Pages:  {analysis['estimated_pages']} (estimated)")
    print("-" * 40)
    
    if analysis['skip_refinement']:
        print(f"   ⚡ FAST MODE: {analysis['reason']}")
    else:
        print("   🧠 FULL MODE: LLM refinement will run")
    print("=" * 40 + "\n")


# =============================================================================
# SECTION 3: ROBUST ERROR HANDLING
# =============================================================================

T = TypeVar('T')


def safe_llm_invoke(
    llm,
    prompt,
    max_retries: int = 3,
    fallback_value: Any = None,
    operation_name: str = "LLM Call"
) -> Any:
    """
    Pattern: Graceful Degradation
    Wraps LLM calls with retry logic and exponential backoff.
    
    Args:
        llm: The LLM instance to invoke
        prompt: The prompt to send
        max_retries: Maximum number of retry attempts
        fallback_value: Value to return if all retries fail
        operation_name: Name for logging purposes
        
    Returns:
        LLM response or fallback_value on failure
    """
    last_error = None
    
    for attempt in range(max_retries):
        try:
            return llm.invoke(prompt)
        except Exception as e:
            last_error = e
            wait_time = 2 ** attempt  # Exponential backoff: 1s, 2s, 4s
            
            print(f"   ⚠️ {operation_name} failed (attempt {attempt + 1}/{max_retries}): {type(e).__name__}")
            
            if attempt < max_retries - 1:
                print(f"      Retrying in {wait_time}s...")
                time.sleep(wait_time)
    
    print(f"   ❌ {operation_name} exhausted retries. Using fallback.")
    
    # If fallback is a callable, call it to get the value
    if callable(fallback_value):
        return fallback_value()
    return fallback_value


def extract_json_robust(text: str) -> str:
    """
    Multiple strategies to extract JSON from LLM output.
    Handles markdown code blocks, nested JSON, etc.
    """
    if not text:
        return "{}"
    
    # Strategy 1: Try to find JSON in markdown code block
    code_block_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
    if code_block_match:
        return code_block_match.group(1).strip()
    
    # Strategy 2: Find JSON object pattern (handles nested braces)
    brace_count = 0
    start_idx = None
    
    for i, char in enumerate(text):
        if char == '{':
            if start_idx is None:
                start_idx = i
            brace_count += 1
        elif char == '}':
            brace_count -= 1
            if brace_count == 0 and start_idx is not None:
                return text[start_idx:i + 1]
    
    # Strategy 3: Simple regex (original fallback)
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        return match.group(0)
    
    return text


def safe_json_parse(
    text: str,
    schema_class: Type[BaseModel],
    fallback_factory: Callable[[], T],
    operation_name: str = "JSON Parse"
) -> T:
    """
    Pattern: Defensive Parsing
    Attempts multiple extraction strategies before using fallback.
    
    Args:
        text: Raw text from LLM that should contain JSON
        schema_class: Pydantic model to validate against
        fallback_factory: Callable that returns fallback instance
        operation_name: Name for logging
        
    Returns:
        Parsed and validated model instance, or fallback
    """
    if not text:
        print(f"   ⚠️ {operation_name}: Empty input. Using fallback.")
        return fallback_factory()
    
    extraction_strategies = [
        ("Direct parse", lambda t: t),
        ("JSON extraction", extract_json_robust),
        ("Strip whitespace", lambda t: extract_json_robust(t.strip())),
    ]
    
    for strategy_name, extractor in extraction_strategies:
        try:
            extracted = extractor(text)
            result = schema_class.model_validate_json(extracted)
            return result
        except Exception as e:
            continue  # Try next strategy
    
    print(f"   ⚠️ {operation_name} failed all strategies. Using fallback.")
    return fallback_factory()


class FallbackResponse:
    """Mock response object for LLM fallback scenarios."""
    def __init__(self, content: str = ""):
        self.content = content


def create_fallback_response(content: str = "Fallback: Analysis unavailable.") -> FallbackResponse:
    """Factory for creating fallback LLM responses."""
    return FallbackResponse(content)


# =============================================================================
# SECTION 4: ENHANCED NODE WRAPPERS
# =============================================================================

def safe_code_execution(sandbox, code: str, max_output_len: int = 1000) -> str:
    """
    Safely execute code with output truncation.
    
    Args:
        sandbox: The sandbox executor instance
        code: Code string to execute
        max_output_len: Maximum output length before truncation
        
    Returns:
        Execution output (possibly truncated)
    """
    try:
        output = sandbox.execute(code)
        
        # Truncate very long outputs
        if len(output) > max_output_len:
            output = output[:max_output_len] + f"\n... [truncated {len(output) - max_output_len} chars]"
        
        return output
    except Exception as e:
        return f"EXECUTION_ERROR: {type(e).__name__}: {str(e)}"


def should_attempt_fix(output: str, failed_fixes: int, max_failures: int = 3) -> bool:
    """
    Determine if we should attempt to fix a code error.
    
    Args:
        output: The execution output
        failed_fixes: Number of fixes already failed
        max_failures: Maximum fix attempts allowed
        
    Returns:
        True if we should attempt a fix
    """
    if failed_fixes >= max_failures:
        return False
    
    # Don't try to fix certain types of errors
    unfixable_patterns = [
        "Security Violation",
        "EXECUTION_ERROR",
        "ModuleNotFoundError",  # Can't fix missing imports
    ]
    
    for pattern in unfixable_patterns:
        if pattern in output:
            return False
    
    return "ERROR" in output or "Traceback" in output


# =============================================================================
# SECTION 5: UTILITY FUNCTIONS
# =============================================================================

def update_thresholds(max_chunks: int = None, max_words: int = None, max_pages: int = None):
    """
    Update the global refinement thresholds.
    
    Args:
        max_chunks: New max chunks threshold (or None to keep current)
        max_words: New max words threshold (or None to keep current)
        max_pages: New max pages threshold (or None to keep current)
    """
    global REFINEMENT_THRESHOLDS
    
    if max_chunks is not None:
        REFINEMENT_THRESHOLDS["max_chunks"] = max_chunks
    if max_words is not None:
        REFINEMENT_THRESHOLDS["max_words"] = max_words
    if max_pages is not None:
        REFINEMENT_THRESHOLDS["max_pages"] = max_pages
    
    print(f"📊 Updated thresholds: chunks={REFINEMENT_THRESHOLDS['max_chunks']}, "
          f"words={REFINEMENT_THRESHOLDS['max_words']:,}, "
          f"pages={REFINEMENT_THRESHOLDS['max_pages']}")


def get_file_stats(file_path: str) -> Dict[str, Any]:
    """
    Get basic statistics about a file before full processing.
    Quick check without loading entire content.
    """
    stats = {
        "exists": os.path.exists(file_path),
        "size_bytes": 0,
        "size_mb": 0.0,
        "extension": "",
    }
    
    if stats["exists"]:
        stats["size_bytes"] = os.path.getsize(file_path)
        stats["size_mb"] = stats["size_bytes"] / (1024 * 1024)
        stats["extension"] = os.path.splitext(file_path)[1].lower()
    
    return stats


# =============================================================================
# EXPORT ALL PUBLIC FUNCTIONS
# =============================================================================

__all__ = [
    # Document specialists
    "TXTSpecialist",
    "MarkdownSpecialist", 
    "DOCXSpecialist",
    "PDFSpecialist",
    "register_all_specialists",
    # Pre-flight
    "REFINEMENT_THRESHOLDS",
    "preflight_analysis",
    "print_preflight_report",
    "update_thresholds",
    # Error handling
    "safe_llm_invoke",
    "safe_json_parse",
    "extract_json_robust",
    "FallbackResponse",
    "create_fallback_response",
    # Code execution
    "safe_code_execution",
    "should_attempt_fix",
    # Utilities
    "get_file_stats",
]
